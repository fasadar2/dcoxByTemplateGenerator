import tkinter.messagebox
from tkinter import ttk, IntVar
import tkinter as tk
import docx
import xlrd
from threading import Thread
import os

def CreateDocThread(data: dict, templatePath: str, coursor: int, keys: list,progressbar:ttk.Progressbar) -> None:
    try:
        doc = docx.Document(templatePath)
        for key in keys:

            for parag in doc.paragraphs:
                if key in parag.text:
                    parag.text = parag.text.replace(key, data[key][coursor])

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if key in cell.text:
                            cell.text = cell.text.replace(key, data[key][coursor])
        progressbar.step()

        doc.save(f"bin/{coursor}.docx")

    except Exception as e:
        tkinter.messagebox.showerror(title="Исключеньеце необработанное", message=str(e))


def CreateDocByTemplate(data: dict, templatePath: str) -> None:
    keys = list(data.keys())
    value_var = IntVar(value=len(data[keys[0]]))
    progressbar = ttk.Progressbar(orient="horizontal",variable=value_var, length=590)
    progressbar.grid(column=1, row=7, columnspan=2, ipadx=6, ipady=6, padx=4, pady=4, sticky=tk.SE)
    try:
        if not os.path.isdir('bin'):
            os.mkdir("bin")
        progressbar.start()
        for i in range(len(data[keys[0]])):
            Thread(target=CreateDocThread, args=(data, templatePath, i, keys, progressbar)).start()
        progressbar.stop()
        tkinter.messagebox.showinfo(title="Вас посетил успех", message="Файлы успешно созданы")
        progressbar.destroy()
    except Exception as e:
        tkinter.messagebox.showerror(title="Исключеньеце необработанное", message=str(e))


def DecompileExcel(dataPath: str) -> dict:
    rezult = {}
    try:
        # Open the Workbook
        workbook = xlrd.open_workbook(dataPath)

        # Выбор листа
        sheet = workbook.sheet_by_index(0)  # Или можно использовать sheet_by_name('имя_листа')

        # Получение первой строки
        first_row = sheet.row_values(0)

        # Формирование списка ключей
        keys_list = [str(key) for key in first_row]
        for key in keys_list:
            rezult[key] = []
        for i in range(1, sheet.nrows):
            row = sheet.row_values(i)
            for ii in range(len(keys_list)):
                rezult[keys_list[ii]].append(row[ii])
    except Exception as e:
        tkinter.messagebox.showerror(title="Исключеньеце необработанное", message=str(e))

    return rezult


def CreateDocx(dataPath: str, templatePath: str):
    data = DecompileExcel(dataPath)
    CreateDocByTemplate(data, templatePath)
